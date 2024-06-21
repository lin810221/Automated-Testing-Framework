# -*- coding: utf-8 -*-
import ast
import os
import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import RGBColor

# =============================================================================
#                                匯出 Excel 測試結果
# =============================================================================
def save_to_excel(df, excel_file, folder_path = 'record', end_datetime = datetime.datetime.now()):
    try:
        print('Saving Excel file...')
        # 檢查文件夾是否存在，若不存在則創建
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        base_name, extension = os.path.splitext(excel_file)
        excel_fname = f'{base_name}_{end_datetime.strftime("%Y-%m-%d-%H-%M-%S")}.xlsx'
        excel_file_path = os.path.join(folder_path, excel_fname)
        df.to_excel(excel_file_path, index=False)
        
        return excel_fname
    except Exception as e:
        print(e)
        return None

# =============================================================================
#                                      QA 內部報告
# =============================================================================
def report_QA(test_plan, excel_file, start_time, end_time = datetime.datetime.now(), folder_path = 'QA_Report'):
    try:
        print('Saving QA Report...')
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        # 匯入 dataframe
        df_Result = test_plan[['用例編號', '測試標題', '測試項目', '測試結果']]
        
        # 計算良率與不良率
        pass_count = df_Result['測試結果'].value_counts().get('Pass', 0)
        fail_count = df_Result['測試結果'].value_counts().get('Fail', 0)
        total_count = pass_count + fail_count
        
        pass_rate = (pass_count/total_count) * 100 if total_count > 0 else 0
        fail_rate = (fail_count/total_count) * 100 if total_count > 0 else 0
        
        title, extension = os.path.splitext(excel_file)
        head = '''
        <head>
            <meta charset="UTF-8">
            <title>QA Report</title>
            <style>
                table {
                    border-collapse: collapse;
                }
                th, td {
                    border: 1px solid black;
                    padding: 5px;
                }
                .pass {
                    background-color: green;
                    color: white;
                }
                .fail {
                    background-color: red;
                    color: white;
                }
                .nan{
                    background-color: orange;
                    color: white;
                    }
            </style>
        </head>
        '''
        
        title_html = f'<h1>{title}</h1>'
        
        rate_html = f'<h3>良率：{pass_rate:.2f}%</h3>\
            <h3>不良率：{fail_rate:.2f}%</h3>'
        
        time_html = f'<p>開始時間：{start_time.strftime("%Y-%m-%d %H:%M:%S")}</p>\
            <p>結束時間：{end_time.strftime("%Y-%m-%d %H:%M:%S")}</p>'
        
        # 將 'Result' 欄位的值轉換為帶有對應 CSS class 的 HTML 標記
        df_Result['測試結果'] = df_Result['測試結果'].apply(lambda x: f'<td class="pass">{x}</td>' if x == 'Pass' else (f'<td class="fail">{x}</td>' if x == 'Fail' else f'<td class="nan">{x}</td>'))
        #df_Result['測試結果'] = df_Result['測試結果'].apply(lambda x: f'<td class="pass">{x}</td>' if x == 'Pass' else f'<td class="fail">{x}</td>')
        df_html = df_Result.to_html(index=False, escape=False)
        df_html = df_html.replace('<td><td class="fail">Fail</td></td>', '<td class="fail">Fail</td>').replace('<td><td class="pass">Pass</td></td>', '<td class="pass">Pass</td>').replace('<td><td class="nan">nan</td></td>', '<td class="nan">nan</td>')
        
        html_content = f'{head}{title_html}{rate_html}{time_html}{df_html}'
        html_fname = f'{title}_{end_time.strftime("%Y-%m-%d-%H-%M-%S")}.html'
        html_file_path = os.path.join(folder_path, html_fname)
        with open(html_file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    except Exception as e:
        print(e)

# =============================================================================
#                               WORD 報告
# =============================================================================
class TestReport:
    def __init__(self):
        print('Saving Word Report...')
        self.doc = Document()

    def create_cover_page(self, df, **kwargs):
        title = self.doc.add_heading('軟體測試報告書', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        result_counts = df['測試結果'].value_counts()
        result_info = ' | '.join([f"{result}: {count}" for result, count in result_counts.items()])

        self.doc.add_heading('摘要資訊', level=1)
        self.doc.add_paragraph(
            f'''
            測試結果：\t{result_info}
            產品名稱：\tNA
            測試版本：\tNA
            設備類型：\t{kwargs.get('device', 'NA')}
            平台類型：\t{kwargs.get('platform', 'NA')}
            測試時間：\t{kwargs.get('test_time', datetime.datetime.now()).strftime("%Y-%m-%d %H:%M:%S")}
            測試人員：\tNA
            單位主管：\tNA
            ''')
        self.doc.add_page_break()

    def create_content_page(self, df):
        self.doc.add_heading('測試紀錄', level = 1)
        
        result_color = {'Pass': RGBColor(0, 128, 0),
                        'Fail': RGBColor(255, 0, 0)}
        
        for index in df.index:
            sn = df.loc[index, '用例編號']
            title = df.loc[index, '測試標題']
            item = df.loc[index, '測試項目']
            result = df.loc[index, '測試結果']
            
            self.doc.add_heading(f'{sn}  {title} - {item}', level = 2)
            
            color = result_color.get(result, RGBColor(255, 140, 0))
            content = self.doc.add_paragraph(result)
            content.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in content.runs:
                run.bold = True
                run.font.color.rgb = color

            dict_data = ast.literal_eval(df.loc[index, '圖片檔名'])

            for key, value in dict_data.items():
                if os.path.exists(value):
                    self.doc.add_paragraph(key)
                    paragraph = self.doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(value, width = Inches(5))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    self.doc.add_paragraph(f"{key}: \n\t找不到圖片 {value}", style = "Normal")
                    
    def save_report(self, excel_file, end_time):
        self.doc.save(f'軟體測試報告_{excel_file}_{end_time.strftime("%Y-%m-%d-%H-%M-%S")}.docx')

